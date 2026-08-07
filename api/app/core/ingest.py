"""Document ingestion pipeline with incremental indexing support.

Implements Workstream A1: Incremental ingestion and indexing.
- Content hashes per document for change detection
- Only re-chunk and re-embed changed documents
- Append to existing index without full rebuild when possible
"""

from __future__ import annotations

import hashlib
import io
import mimetypes
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from threading import Lock
import re
import shutil
import subprocess

try:  # pragma: no cover - optional dependency
    from pypdf import PdfReader  # type: ignore
    print("PDF library (pypdf) loaded successfully")
except ImportError:  # pragma: no cover
    PdfReader = None  # type: ignore
    print("WARNING: PDF library (pypdf) failed to load")

try:  # pragma: no cover
    import docx  # type: ignore
    print("DOCX library loaded successfully")
except ImportError:  # pragma: no cover
    docx = None  # type: ignore
    print("WARNING: docx library failed to load")

from ..schemas.config import AppConfig, OCRPolicy, ProviderConfig
from .audit import AuditAction, AuditEntry, get_audit_log
from .chunking import Chunk
from .contextual_enrichment import ContextualEnricher, EnrichmentConfig
from .documents import DocumentStore
from .document_parser import DocumentParserRouter, parser_result_from_text, parser_result_to_metadata
from .langextract_enricher import LangExtractEnricher
from .local_first import LocalFirstRegistry
from .ocr import OCRRouter
from .prompt_guard import ThreatLevel, sanitize_at_ingest
from .retrieval import RetrievalEngine

_INDEX_EXECUTOR = ThreadPoolExecutor(max_workers=1, thread_name_prefix="autorag-index")
_INDEX_LOCK = Lock()
_LANGEXTRACT_SYNTHETIC_BEGIN = "[[LANGEXTRACT_SYNTHETIC_FACTS_BEGIN]]"
_LANGEXTRACT_SYNTHETIC_END = "[[LANGEXTRACT_SYNTHETIC_FACTS_END]]"


@dataclass
class IngestResult:
    """Result of document ingestion."""
    document_id: str
    title: str
    chunk_count: int
    was_modified: bool = True  # False if skipped due to unchanged content
    content_hash: str = ""  # SHA-256 hash of content for change tracking


class IngestPipeline:
    """Handles text/file ingestion with incremental indexing support.

    Key features (A1 requirement):
    - Content hashing for change detection
    - Skip re-processing of unchanged documents
    - Contextualize chunks with document metadata
    """

    def __init__(
        self,
        store: DocumentStore,
        retrieval: RetrievalEngine,
        config_getter: Callable[[], AppConfig] | None = None,
        data_dir: Path | None = None,
        policy_registry: LocalFirstRegistry | None = None,
    ) -> None:
        self._store = store
        self._retrieval = retrieval
        self._config_getter = config_getter
        self._langextract = LangExtractEnricher(data_dir=data_dir)
        self._policy_registry = policy_registry
        self._parser_router = DocumentParserRouter()
        self._enricher = ContextualEnricher(EnrichmentConfig(
            add_document_title=True,
            add_section_header=True,
            add_chunk_summary=True,
            add_context_window=True,
            use_llm_for_summary=False,  # Sync path uses heuristic summaries
            fallback_to_heuristic=True,
        ))

    def ingest_text(
        self,
        title: str,
        text: str,
        metadata: dict[str, str] | None = None,
        sync: bool = False,
        langextract_profile_override: str | None = None,
        langextract_prompt_override: str | None = None,
        on_duplicate: str = "reject",
    ) -> IngestResult:
        """Ingest text with content hash tracking for change detection."""
        meta = self._prepare_metadata(metadata)
        meta.setdefault("processing_status", "processing")

        sanitized_text, injection_attempts = sanitize_at_ingest(
            text,
            source=f"document:{title}",
            wrap_delimiters=False,
        )
        if injection_attempts:
            threat_order = [ThreatLevel.NONE, ThreatLevel.LOW, ThreatLevel.MEDIUM, ThreatLevel.HIGH, ThreatLevel.CRITICAL]
            highest_threat = max(injection_attempts, key=lambda attempt: threat_order.index(attempt.threat_level)).threat_level
            meta["prompt_injection_detected"] = "true"
            meta["prompt_injection_attempts"] = str(len(injection_attempts))
            meta["prompt_injection_threat_level"] = highest_threat.value

        langextract_result = self._run_langextract(
            text=sanitized_text,
            profile_override=langextract_profile_override,
            prompt_override=langextract_prompt_override,
        )
        augmented_text = self._append_langextract_sections(sanitized_text, langextract_result.get("synthetic_sections", []))
        self._apply_langextract_metadata(meta, langextract_result)

        # Compute content hash for change detection (A1)
        content_hash = self._compute_content_hash(augmented_text)
        meta["content_hash"] = content_hash

        chunks = self._chunk(augmented_text)
        # Scan for poisoned chunks (OWASP LLM02: knowledge base poisoning)
        poison_flags = self._scan_for_poisoned_chunks(chunks)
        if poison_flags:
            meta["poison_warnings"] = str(len(poison_flags))
            meta["poison_flags"] = "; ".join(poison_flags[:5])
        # Contextual enrichment: add document title, section header, summary,
        # and context window to each chunk (Anthropic Contextual Retrieval).
        # Falls back to simple header prepend if enrichment fails.
        contextualized = self._enrich_chunks(chunks, augmented_text, title, meta)
        combined = "\n\n".join(contextualized)

        doc = self._store.add(title=title, text=combined, metadata=meta, on_duplicate=on_duplicate)

        self._persist_langextract_artifact(doc.id, title, langextract_result, doc)
        self._log_langextract_audit(doc.id, title, langextract_result)

        # Run indexing in a shared background worker to keep API responsive
        def do_build() -> None:
            try:
                with _INDEX_LOCK:
                    if hasattr(self._retrieval, "index_documents"):
                        self._retrieval.index_documents([doc])
                    else:
                        self._retrieval.build()
                        if hasattr(self._retrieval, "increment_corpus_version"):
                            self._retrieval.increment_corpus_version()
                doc.metadata["processing_status"] = "ready"
                doc.metadata["processed_at"] = datetime.now(UTC).isoformat()
                self._store.upsert(doc)
            except Exception as exc:
                doc.metadata["processing_status"] = "error"
                doc.metadata["processing_error"] = str(exc)
                self._store.upsert(doc)

        if sync:
            do_build()
        else:
            _INDEX_EXECUTOR.submit(do_build)

        return IngestResult(
            document_id=doc.id,
            title=doc.title,
            chunk_count=len(chunks),
            was_modified=True,
            content_hash=content_hash,
        )

    def ingest_file(
        self,
        title: str,
        content: bytes,
        metadata: dict[str, str] | None = None,
        sync: bool = False,
        langextract_profile_override: str | None = None,
        langextract_prompt_override: str | None = None,
        on_duplicate: str = "reject",
    ) -> IngestResult:
        meta = {**(metadata or {})}
        meta.setdefault("filename", title)
        meta.setdefault("original_filename", meta["filename"])
        meta.setdefault("content_type", mimetypes.guess_type(meta["filename"])[0] or "text/plain")
        meta["filesize"] = str(len(content))
        text, extraction_metadata = self._extract_text_with_metadata(content, meta)
        meta.update(extraction_metadata)
        if "parser_preview_json" not in meta:
            extraction_confidence = self._metadata_float(extraction_metadata, "extraction_confidence")
            if extraction_confidence <= 0:
                extraction_confidence = self._metadata_float(extraction_metadata, "text_extraction_confidence")
            try:
                parser_result = self._parser_router.parse(content, meta)
                if (
                    parser_result.provider == "docling"
                    and parser_result.text.strip()
                    and parser_result.confidence >= extraction_confidence
                ):
                    text = parser_result.text
                    preview_result = parser_result
                else:
                    preview_result = parser_result_from_text(
                        text=text,
                        provider=parser_result.provider,
                        engine=extraction_metadata.get("extraction_engine", parser_result.engine),
                        confidence=max(extraction_confidence, parser_result.confidence),
                        used_ocr=str(extraction_metadata.get("ocr_used", "")).lower() == "true",
                        raw_metadata={"parser_warning": ",".join(parser_result.warnings)},
                    )
                meta.update(parser_result_to_metadata(preview_result, extraction_metadata))
            except Exception as exc:
                meta["parser_provider"] = meta.get("parser_provider", "native")
                meta["parser_warning"] = str(exc)
                preview_result = parser_result_from_text(
                    text=text,
                    provider="native",
                    engine=extraction_metadata.get("extraction_engine", "stored-text"),
                    confidence=max(extraction_confidence, 0.0),
                    used_ocr=str(extraction_metadata.get("ocr_used", "")).lower() == "true",
                    raw_metadata={"parser_warning": str(exc)},
                )
                meta.update(parser_result_to_metadata(preview_result, extraction_metadata))
        return self.ingest_text(
            title=title,
            text=text,
            metadata=meta,
            sync=sync,
            langextract_profile_override=langextract_profile_override,
            langextract_prompt_override=langextract_prompt_override,
            on_duplicate=on_duplicate,
        )

    def ingest_incremental(
        self,
        title: str,
        text: str,
        metadata: dict[str, str] | None = None,
    ) -> IngestResult:
        """Incremental ingest: skip if document content unchanged (A1).

        Args:
            title: Document title (used as identifier)
            text: Document text content
            metadata: Optional metadata dict

        Returns:
            IngestResult with was_modified=False if skipped
        """
        content_hash = self._compute_content_hash(text)

        # Check if document exists with same hash
        existing = self._store.get_by_title(title)
        if existing:
            existing_hash = existing.metadata.get("content_hash", "")
            if existing_hash == content_hash:
                # Document unchanged - skip re-processing
                return IngestResult(
                    document_id=existing.id,
                    title=existing.title,
                    chunk_count=0,
                    was_modified=False,
                    content_hash=content_hash,
                )

        # Proceed with full ingest and intentionally replace the existing title when changed.
        return self.ingest_text(title=title, text=text, metadata=metadata, on_duplicate="replace")

    def _run_langextract(
        self,
        text: str,
        profile_override: str | None,
        prompt_override: str | None,
    ) -> dict[str, object]:
        cfg = self._read_config()
        if cfg is None:
            return {
                "status": "disabled",
                "profile": "",
                "model_source": "",
                "entities_count": 0,
                "relations_count": 0,
                "claims_count": 0,
                "warnings_count": 0,
                "synthetic_sections": [],
                "error": None,
                "raw": None,
            }

        override_bundle = {
            "langextract_profile_override": profile_override,
            "langextract_prompt_override": prompt_override,
        }
        if not self._langextract.is_enabled(cfg, per_doc_override=override_bundle):
            return {
                "status": "disabled",
                "profile": "",
                "model_source": getattr(cfg.retrieval, "langextract_model_source", "gatherer"),
                "entities_count": 0,
                "relations_count": 0,
                "claims_count": 0,
                "warnings_count": 0,
                "synthetic_sections": [],
                "error": None,
                "raw": None,
            }

        profile = (profile_override or "").strip() or getattr(
            cfg.retrieval,
            "langextract_profile_default",
            "generic_entities_v1",
        )
        model_source = getattr(cfg.retrieval, "langextract_model_source", "gatherer")
        timeout = int(getattr(cfg.retrieval, "langextract_timeout_sec", 20))
        max_chars = int(getattr(cfg.retrieval, "langextract_max_chars", 12000))
        max_facts = int(getattr(cfg.retrieval, "langextract_max_synthetic_facts", 200))

        return self._langextract.extract(
            text=text,
            provider=cfg.provider,
            profile=profile,
            prompt_override=prompt_override,
            timeout=timeout,
            model_source=model_source,
            max_chars=max_chars,
            max_synthetic_facts=max_facts,
        )

    def _append_langextract_sections(self, text: str, sections: list[str]) -> str:
        if not sections:
            return text
        body = "\n\n".join(section for section in sections if section.strip())
        if not body:
            return text
        clean_text = text.rstrip()
        return (
            f"{clean_text}\n\n{_LANGEXTRACT_SYNTHETIC_BEGIN}\n"
            f"{body}\n{_LANGEXTRACT_SYNTHETIC_END}\n"
        )

    def _apply_langextract_metadata(
        self,
        metadata: dict[str, str],
        result: dict[str, object],
    ) -> None:
        metadata["langextract_status"] = str(result.get("status", "disabled"))
        metadata["langextract_profile"] = str(result.get("profile", ""))
        metadata["langextract_model_source"] = str(result.get("model_source", ""))
        metadata["langextract_entities_count"] = str(result.get("entities_count", 0))
        metadata["langextract_relations_count"] = str(result.get("relations_count", 0))
        metadata["langextract_claims_count"] = str(result.get("claims_count", 0))
        metadata["langextract_warnings_count"] = str(result.get("warnings_count", 0))
        error = str(result.get("error", "") or "").strip()
        if error:
            metadata["langextract_error"] = error
        else:
            metadata.pop("langextract_error", None)

    def _persist_langextract_artifact(
        self,
        doc_id: str,
        title: str,
        result: dict[str, object],
        doc,
    ) -> None:
        status = str(result.get("status", "disabled"))
        if status == "disabled":
            return
        payload = {
            "document_id": doc_id,
            "title": title,
            "created_at": datetime.now(UTC).isoformat(),
            "status": status,
            "profile": result.get("profile"),
            "model_source": result.get("model_source"),
            "model_id": result.get("model_id"),
            "provider": result.get("provider"),
            "entities_count": result.get("entities_count", 0),
            "relations_count": result.get("relations_count", 0),
            "claims_count": result.get("claims_count", 0),
            "warnings_count": result.get("warnings_count", 0),
            "error": result.get("error"),
            "synthetic_sections": result.get("synthetic_sections", []),
            "raw": result.get("raw"),
        }
        try:
            artifact = self._langextract.persist_artifact(doc_id, payload)
            doc.metadata["langextract_artifact_path"] = str(artifact)
            self._store.upsert(doc)
        except Exception as exc:
            doc.metadata["langextract_error"] = str(exc)
            self._store.upsert(doc)

    def _log_langextract_audit(self, doc_id: str, title: str, result: dict[str, object]) -> None:
        status = str(result.get("status", "disabled"))
        details = {
            "document_id": doc_id,
            "title": title,
            "langextract_status": status,
            "langextract_profile": str(result.get("profile", "")),
            "langextract_model_source": str(result.get("model_source", "")),
            "langextract_entities_count": int(result.get("entities_count", 0) or 0),
            "langextract_relations_count": int(result.get("relations_count", 0) or 0),
            "langextract_claims_count": int(result.get("claims_count", 0) or 0),
            "langextract_warnings_count": int(result.get("warnings_count", 0) or 0),
            "langextract_enabled": status != "disabled",
        }
        get_audit_log().log(
            AuditEntry(
                timestamp=datetime.now(UTC),
                action=AuditAction.INGEST,
                details=details,
                success=True,
            )
        )

    def _read_config(self) -> AppConfig | None:
        if self._config_getter is None:
            return None
        try:
            return self._config_getter()
        except Exception:
            return None

    def _compute_content_hash(self, text: str) -> str:
        """Compute SHA-256 hash of document content for change detection."""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()

    def _metadata_float(self, metadata: dict[str, str], key: str) -> float:
        try:
            return float(metadata.get(key, "0") or 0)
        except ValueError:
            return 0.0

    def _scan_for_poisoned_chunks(self, chunks: list[str]) -> list[str]:
        """Scan chunks for poisoning indicators (OWASP LLM02).

        Returns a list of warning strings for suspicious chunks.
        """
        try:
            from .prompt_guard import get_poison_scanner
            scanner = get_poison_scanner()
            chunk_tuples = [(str(i), text) for i, text in enumerate(chunks)]
            results = scanner.scan_chunks(chunk_tuples)
            warnings = []
            for r in results:
                if r.is_suspicious:
                    warnings.append(f"chunk {r.chunk_id}: {', '.join(r.flags)}")
            return warnings
        except Exception:
            return []

    def _enrich_chunks(
        self,
        chunks: list[str],
        document_text: str,
        doc_title: str,
        metadata: dict[str, str] | None,
    ) -> list[str]:
        """Enrich chunks with contextual retrieval (Anthropic approach).

        Uses ContextualEnricher to add document title, section header,
        chunk summary, and context window to each chunk. Falls back to
        the simple header prepend if enrichment fails.
        """
        try:
            # Build Chunk objects with positions for the enricher
            chunk_objs: list[Chunk] = []
            char_pos = 0
            for i, text in enumerate(chunks):
                stripped = text.strip()
                if stripped:
                    chunk_objs.append(Chunk(
                        text=stripped,
                        index=i,
                        start_char=char_pos,
                        end_char=char_pos + len(stripped),
                    ))
                char_pos += len(text) + 2  # approximate gap

            enriched = self._enricher.enrich_chunks_sync(
                chunks=chunk_objs,
                document_text=document_text,
                filename=metadata.get("filename", doc_title) if metadata else doc_title,
            )
            return [e.enriched_text for e in enriched]
        except Exception:
            return self._contextualize_chunks(chunks, doc_title, metadata)

    def _contextualize_chunks(
        self,
        chunks: list[str],
        doc_title: str,
        metadata: dict[str, str] | None,
    ) -> list[str]:
        """Add header context to each chunk for better interpretability (D2).

        Fallback when ContextualEnricher is unavailable.
        """
        header_parts = [f"[Document: {doc_title}]"]
        if metadata:
            if "date" in metadata:
                header_parts.append(f"[Date: {metadata['date']}]")
            elif "uploaded_at" in metadata:
                # Use upload date if no explicit date
                upload_date = metadata["uploaded_at"][:10]  # YYYY-MM-DD
                header_parts.append(f"[Date: {upload_date}]")
            if "author" in metadata:
                header_parts.append(f"[Author: {metadata['author']}]")

        header = " ".join(header_parts)
        return [f"{header}\n{chunk}" for chunk in chunks]

    def _prepare_metadata(self, metadata: dict[str, str] | None) -> dict[str, str]:
        meta = {**(metadata or {})}
        meta.setdefault("uploaded_at", datetime.now(UTC).isoformat())
        return meta

    def _infer_extension(self, metadata: dict[str, str] | None) -> str:
        if metadata:
            filename = metadata.get("filename")
            if filename:
                return Path(filename).suffix.lower()
            content_type = metadata.get("content_type")
            if content_type:
                return (mimetypes.guess_extension(content_type) or "").lower()
        return ""

    def _detect_magic_extension(self, content: bytes) -> str:
        if content.startswith(b"%PDF"):
            return ".pdf"
        if content[:4] == b"PK\x03\x04":
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as archive:
                    if "word/document.xml" in archive.namelist():
                        return ".docx"
            except Exception:
                return ""
        if content[:8] == b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1":
            return ".doc"
        return ""

    def _extract_text(self, content: bytes, metadata: dict[str, str] | None = None) -> str:
        text, _ = self._extract_text_with_metadata(content, metadata)
        return text

    def _extract_text_with_metadata(
        self,
        content: bytes,
        metadata: dict[str, str] | None = None,
    ) -> tuple[str, dict[str, str]]:
        if self._policy_registry is not None:
            self._policy_registry.ensure_runtime_allowed("document_parser")
        ext = self._infer_extension(metadata)
        if not ext:
            ext = self._detect_magic_extension(content)
        if ext in {".md", ".markdown"}:
            text = self._extract_markdown(content)
            return text, self._build_extraction_metadata(
                extraction_method="markdown_decode",
                extraction_engine="markdown",
                extraction_confidence=1.0 if text.strip() else 0.0,
                ocr_policy=self._resolve_ocr_policy(metadata).value,
            )
        if ext == ".pdf":
            return self._extract_pdf_with_routing(content, metadata)
        if ext in {".doc", ".docx"}:
            text = ""
            method = "docx_native" if ext == ".docx" else "doc_legacy"
            if ext == ".docx":
                text = self._extract_docx_text(content)
            if ext == ".doc":
                print("Tip: .doc files are legacy. Convert to .docx for better extraction.")
            if text.strip():
                return text, self._build_extraction_metadata(
                    extraction_method=method,
                    extraction_engine="python_docx",
                    extraction_confidence=0.95,
                    ocr_policy=self._resolve_ocr_policy(metadata).value,
                )

        try:
            text = content.decode("utf-8")
            return text, self._build_extraction_metadata(
                extraction_method="plain_text_decode",
                extraction_engine="utf8",
                extraction_confidence=1.0 if text.strip() else 0.0,
                ocr_policy=self._resolve_ocr_policy(metadata).value,
            )
        except UnicodeDecodeError:
            text = f"(Binary content: {ext or 'unknown'})"
            return text, self._build_extraction_metadata(
                extraction_method="binary_unsupported",
                extraction_engine="binary",
                extraction_confidence=0.0,
                ocr_policy=self._resolve_ocr_policy(metadata).value,
            )

    def _extract_pdf_with_routing(
        self,
        content: bytes,
        metadata: dict[str, str] | None = None,
    ) -> tuple[str, dict[str, str]]:
        native_text = self._extract_pdf_text(content)
        native_confidence = self._estimate_text_confidence(native_text)

        pdftotext_text = self._extract_pdf_text_pdftotext(content)
        pdftotext_confidence = self._estimate_text_confidence(pdftotext_text)

        extracted_text = native_text
        extraction_engine = "pypdf"
        extraction_method = "native_text"
        extraction_confidence = native_confidence

        if pdftotext_confidence > extraction_confidence:
            extracted_text = pdftotext_text
            extraction_engine = "pdftotext"
            extraction_method = "layout_text"
            extraction_confidence = pdftotext_confidence

        ocr_policy = self._resolve_ocr_policy(metadata)
        if ocr_policy != OCRPolicy.OFF and self._policy_registry is not None:
            self._policy_registry.ensure_runtime_allowed("ocr")

        router = OCRRouter(
            self._read_ocr_settings(metadata),
            provider_config=self._read_provider_config(),
            vision_model=self._read_vision_model(),
            vision_max_pages=self._read_vision_max_pages(),
        )
        routed = router.route(extracted_text, content)

        final_text = routed.text.strip() or extracted_text.strip()
        final_method = routed.method or extraction_method
        final_engine = routed.engine or extraction_engine
        final_confidence = max(routed.confidence, extraction_confidence if final_text == extracted_text.strip() else 0.0)

        if not final_text:
            final_text = (
                "(PDF extraction failed: no text found. "
                "If this is a scanned PDF, install poppler + tesseract for OCR. "
                "If it is text-based, try re-exporting the PDF.)"
            )
            final_method = "pdf_failed"
            final_engine = "none"
            final_confidence = 0.0

        return final_text, self._build_extraction_metadata(
            extraction_method=final_method,
            extraction_engine=final_engine,
            extraction_confidence=final_confidence,
            ocr_policy=ocr_policy.value,
            text_extraction_confidence=extraction_confidence,
            ocr_used=routed.used_ocr,
            ocr_attempted=",".join(routed.attempted),
        )

    def _extract_markdown(self, content: bytes) -> str:
        text = content.decode("utf-8", errors="ignore")
        # lightweight removal of common markdown tokens
        replacements = ["#", "*", "`", ">", "- ", "* "]
        for token in replacements:
            text = text.replace(token, "")
        return text

    def _extract_pdf_text(self, content: bytes) -> str:
        if not PdfReader:
            return ""
        try:
            reader = PdfReader(io.BytesIO(content), strict=False)  # type: ignore[name-defined]
            if getattr(reader, "is_encrypted", False):
                try:
                    reader.decrypt("")  # type: ignore[attr-defined]
                except Exception:
                    return ""
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except Exception as exc:
            print(f"Error extracting PDF text: {exc}")
            return ""

    def _extract_pdf_text_pdftotext(self, content: bytes) -> str:
        """Fallback PDF extraction using pdftotext if installed."""
        try:
            pdftotext_bin = shutil.which("pdftotext")
            if not pdftotext_bin:
                return ""
            with tempfile.TemporaryDirectory() as tmpdir:
                pdf_path = Path(tmpdir) / "input.pdf"
                txt_path = Path(tmpdir) / "output.txt"
                pdf_path.write_bytes(content)
                result = subprocess.run(
                    [pdftotext_bin, "-enc", "UTF-8", str(pdf_path), str(txt_path)],
                    check=False,
                    capture_output=True,
                    text=True,
                )
                if result.returncode != 0:
                    return ""
                if not txt_path.exists():
                    return ""
                return txt_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from a .docx using python-docx or a zip/XML fallback."""
        if docx:
            try:
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
                if text.strip():
                    return text
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        table_text.append(" | ".join([cell.text for cell in row.cells]))
                return "\n".join(table_text)
            except Exception as exc:
                print(f"Error extracting Word text: {exc}")

        # Fallback: parse the document.xml from the docx zip
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                xml_data = archive.read("word/document.xml")
            root = ET.fromstring(xml_data)
            namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            lines: list[str] = []
            buffer: list[str] = []
            for elem in root.iter():
                if elem.tag == f"{namespace}t" and elem.text:
                    buffer.append(elem.text)
                elif elem.tag == f"{namespace}p" and buffer:
                    lines.append("".join(buffer).strip())
                    buffer = []
            if buffer:
                lines.append("".join(buffer).strip())
            return "\n".join([line for line in lines if line])
        except Exception as exc:
            print(f"Error parsing DOCX fallback: {exc}")
            return ""

    def _chunk(self, text: str, target: int = 800) -> list[str]:
        cfg = self._read_config()
        strategy = "fixed"
        overlap = 50
        if cfg is not None:
            target = max(200, int(getattr(cfg.retrieval, "chunk_size", target)))
            strategy = str(getattr(cfg.retrieval, "chunking_strategy", "fixed")).lower()
            overlap = max(0, int(getattr(cfg.retrieval, "chunk_overlap", overlap)))

        # Route to the shared chunking module for non-fixed strategies so the
        # configured chunking_strategy actually takes effect. The inline
        # header-aware splitter below remains the "fixed" default for backward
        # compatibility.
        if strategy in {"semantic", "recursive", "late"}:
            try:
                from .chunking import ChunkingStrategy, get_chunker

                chunker = get_chunker(
                    strategy=ChunkingStrategy(strategy),
                    target_size=target,
                    overlap=overlap,
                )
                chunk_objs = chunker.chunk(text)
                texts = [c.text for c in chunk_objs if c.text and c.text.strip()]
                return texts or [text.strip()]
            except Exception:
                # Fall back to the inline splitter if the chunking module or an
                # optional dependency (e.g. sentence-transformers for semantic)
                # is unavailable.
                pass

        clean = text.replace("\r", "").replace("\f", "\n\n[PAGE_BREAK]\n\n")
        raw_blocks = [block.strip() for block in re.split(r"\n{2,}", clean) if block.strip()]

        sections: list[str] = []
        pending: list[str] = []
        for block in raw_blocks:
            is_header = bool(re.match(r"^(#{1,6}\s+.+|[A-Z][A-Z0-9\s:/-]{4,}|Section\s+\d+[:.]?.*)$", block))
            if is_header and pending:
                sections.append("\n\n".join(pending))
                pending = [block]
                continue
            pending.append(block)
        if pending:
            sections.append("\n\n".join(pending))

        chunks: list[str] = []
        current = ""
        overlap = max(0, min(overlap or target // 8, target // 2))
        for section in sections:
            candidate = f"{current}\n\n{section}".strip() if current else section
            if current and len(candidate) > target:
                chunks.append(current.strip())
                current = current[-overlap:].strip()
                candidate = f"{current}\n\n{section}".strip() if current else section
            if len(section) > target * 1.25:
                paragraphs = [p.strip() for p in section.split("\n\n") if p.strip()]
                for para in paragraphs:
                    candidate = f"{current}\n\n{para}".strip() if current else para
                    if current and len(candidate) > target:
                        chunks.append(current.strip())
                        current = current[-overlap:].strip()
                    current = f"{current}\n\n{para}".strip() if current else para
            else:
                current = candidate
        if current.strip():
            chunks.append(current.strip())
        return chunks or [text.strip()]

    def _estimate_text_confidence(self, text: str) -> float:
        stripped = text.strip()
        if not stripped:
            return 0.0
        alpha = sum(1 for ch in stripped if ch.isalpha())
        printable = sum(1 for ch in stripped if ch.isprintable() and not ch.isspace())
        density = min(len(stripped) / 600.0, 1.0)
        alpha_ratio = alpha / max(printable, 1)
        return max(0.0, min((density * 0.55) + (alpha_ratio * 0.45), 1.0))

    def _resolve_ocr_policy(self, metadata: dict[str, str] | None = None) -> OCRPolicy:
        override = (metadata or {}).get("ocr_policy", "").strip().lower()
        if override:
            try:
                return OCRPolicy(override)
            except ValueError:
                pass
        cfg = self._read_config()
        if cfg is None:
            return OCRPolicy.AUTO
        return getattr(cfg.ingest.ocr, "policy", OCRPolicy.AUTO)

    def _read_ocr_settings(self, metadata: dict[str, str] | None = None):
        cfg = self._read_config()
        if cfg is None:
            return type("FallbackOCRSettings", (), {
                "policy": self._resolve_ocr_policy(metadata),
                "extractable_text_threshold": 0.65,
                "min_characters": 80,
                "allow_cloud_fallback": False,
                "preferred_backends": ["ocr.local.tesseract", "ocr.local.vision"],
                "dual_merge_strategy": "highest_confidence",
            })()
        settings = cfg.ingest.ocr.model_copy()
        policy = self._resolve_ocr_policy(metadata)
        if policy != settings.policy:
            settings.policy = policy
        return settings

    def _read_provider_config(self) -> ProviderConfig | None:
        cfg = self._read_config()
        if cfg is None:
            return None
        return cfg.provider

    def _read_vision_model(self) -> str | None:
        cfg = self._read_config()
        if cfg is None:
            return None
        ocr_backend = cfg.backends.get("ocr")
        if ocr_backend is not None:
            explicit_model = ocr_backend.settings.get("vision_model")
            if isinstance(explicit_model, str) and explicit_model.strip():
                return explicit_model.strip()
        provider = cfg.provider
        if provider is None:
            return None
        return provider.generator_model or provider.gatherer_model or provider.planner_model

    def _read_vision_max_pages(self) -> int:
        cfg = self._read_config()
        if cfg is None:
            return 8
        ocr_backend = cfg.backends.get("ocr")
        if ocr_backend is None:
            return 8
        raw = ocr_backend.settings.get("max_pages")
        if isinstance(raw, int) and raw > 0:
            return raw
        if isinstance(raw, str) and raw.isdigit():
            return max(1, int(raw))
        return 8

    def _build_extraction_metadata(
        self,
        *,
        extraction_method: str,
        extraction_engine: str,
        extraction_confidence: float,
        ocr_policy: str,
        text_extraction_confidence: float | None = None,
        ocr_used: bool = False,
        ocr_attempted: str = "",
    ) -> dict[str, str]:
        details = {
            "extraction_method": extraction_method,
            "extraction_engine": extraction_engine,
            "extraction_confidence": f"{extraction_confidence:.3f}",
            "ocr_policy": ocr_policy,
            "ocr_used": "true" if ocr_used else "false",
            "last_indexed_at": datetime.now(UTC).isoformat(),
        }
        if text_extraction_confidence is not None:
            details["text_extraction_confidence"] = f"{text_extraction_confidence:.3f}"
        if ocr_attempted:
            details["ocr_attempted"] = ocr_attempted
        return details
